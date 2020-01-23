#pip install pthon-docx 
import docx,re
import tkinter.filedialog as fd
from tkinter import*
#d=docx.Document('H:\\ats practice\\demo.docx')
inputfile = fd.askopenfilename(initialdir = "/",title = "Select file",
                    filetypes = (("doc files","*.docx"),("all files","*.*")))
print(inputfile)
d=docx.Document(inputfile)
leng=len(d.paragraphs)
#print(docx.section[0].text())
for i in range(leng):
     text=d.paragraphs[i].text
     if text.lower() not in [r'resume[a-z]*',r'personal[a-z]*',r'cv|c[a-z]*']:
         if text.lower().startswith('name:')or text.lower().startswith('name-'):
             text=re.sub(r'name:|name-','',text.lower())
         name=re.search(r'[a-zA-Z]+[ .A-Za-z]*',text).group()
         print(name)
         break

for i in range(leng):
    text=d.paragraphs[i].text
    mail=re.findall(r'[a-z-_]+[a-z0-9-_]*@[a-z-_]+.com',text)
    if mail!=[]:
        break
for i in range(leng):
    text=d.paragraphs[i].text
    phno=re.findall(r'[0|+91|91]\d{10,13}',text)
    if phno!=[]:
        break
print(mail)
l=['JAVA','ASP.NET','C#','C++','IOSDEVELOPER','ANDROID DEVELOPER',
   'PYTHON','HTML','CSS','RUBY','FLASK','DJANGO'
   ,'JAVASCRIPT','NODEJS','HADOOP','DATASCIENCE','CYBER SECURITY']
'''for i in range(5,25):
     text=d.paragraphs[i].text
     for word in text:
          if word in l:
               print(i)
               print(word,end="\t")'''
skills=[]
for i in range(5,25):
     text=d.paragraphs[i].text
     #qua=re.search(r'\w*\s*btech|mtech|mba|mca\s*\w*',text.lower()).group()
     for j in l:
          if j in text.upper():
               print(j)
               skills.append(j)
for i in range(1,30):
     text=d.paragraphs[i].text
     exp=re.findall(r'@\W*\d+years$|@\d+\W*yr$',text.lower())

print(phno,skills,exp)


